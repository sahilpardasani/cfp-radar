from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "CFP_Radar_Product_and_Systems_Guide.docx"

BLUE = RGBColor(31, 78, 121)
TEAL = RGBColor(27, 116, 126)
DARK = RGBColor(34, 45, 58)
MUTED = RGBColor(92, 103, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_TEAL = "E8F3F4"
LIGHT_GRAY = "F2F4F7"
WHITE = RGBColor(255, 255, 255)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_numbering_definition(doc, num_id, bullet=False):
    numbering = doc.part.numbering_part.element
    abstract_id = max([int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))] or [0]) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    lvl.append(p_pr)
    if bullet:
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(fonts)
        lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    numbering.append(num)


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def style_run(run, size=11, bold=False, italic=False, color=DARK, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_body(doc, text, bold_lead=None, after=6, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_together = keep
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        style_run(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        style_run(rest)
    else:
        style_run(p.add_run(text))
    return p


def add_bullet(doc, text, num_id=41):
    p = doc.add_paragraph()
    apply_num(p, num_id)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    style_run(p.add_run(text))
    return p


def add_step(doc, title, detail, num_id=42):
    p = doc.add_paragraph()
    apply_num(p, num_id)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.25
    style_run(p.add_run(title + " — "), bold=True, color=BLUE)
    style_run(p.add_run(detail))
    return p


def add_callout(doc, label, text, fill=LIGHT_TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.keep_together = True
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "1B747E")
    borders.append(left)
    p_pr.append(borders)
    style_run(p.add_run(label + ": "), bold=True, color=TEAL)
    style_run(p.add_run(text), color=DARK)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.keep_together = True
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F5F7F9")
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "5")
        border.set(qn("w:color"), "D7DCE2")
        borders.append(border)
    p_pr.append(borders)
    lines = text.splitlines()
    for index, line in enumerate(lines):
        if index:
            p.add_run().add_break()
        style_run(p.add_run(line), size=9.25, color=DARK, font="Courier New")
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    style_run(p.add_run(text), size={1: 16, 2: 13, 3: 12}[level], bold=True,
              color=BLUE if level < 3 else TEAL)
    return p


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        style_run(p.add_run(header), size=9.5, bold=True, color=BLUE)
    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            style_run(p.add_run(value), size=9.25, color=DARK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, before, after, color in [
        (1, 16, 18, 9, BLUE), (2, 13, 14, 7, BLUE), (3, 12, 10, 5, TEAL)
    ]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    style_run(header.add_run("CFP RADAR  /  PRODUCT & SYSTEMS GUIDE"), size=8.5, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    add_numbering_definition(doc, 41, bullet=True)
    add_numbering_definition(doc, 42, bullet=False)
    add_numbering_definition(doc, 43, bullet=False)
    add_numbering_definition(doc, 44, bullet=False)
    add_numbering_definition(doc, 45, bullet=False)


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    # Cover
    doc.add_paragraph().paragraph_format.space_after = Pt(72)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(kicker.add_run("PRODUCT & SYSTEMS GUIDE"), size=11, bold=True, color=TEAL)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(10)
    style_run(title.add_run("How CFP Radar Works"), size=30, bold=True, color=DARK)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    style_run(subtitle.add_run("Product architecture, data workflow, technical choices, trust model, and tradeoffs"), size=15, color=BLUE)
    audience = doc.add_paragraph()
    audience.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(audience.add_run("Written for a CS-trained product builder who wants system-level clarity without implementation-heavy detail"), size=11.5, italic=True, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(72)
    stamp = doc.add_paragraph()
    stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(stamp.add_run("Current system: July 2026  |  Next.js 14  |  Automated every two days"), size=10.5, bold=True, color=TEAL)
    doc.add_page_break()

    add_heading(doc, "How this guide is structured", 1)
    add_body(doc, "The guide starts with the product surface and system boundaries, then follows data from discovery through admission, persistence, APIs, and presentation. Later sections cover the stack, operating model, scaling boundary, and architectural alternatives.")
    add_callout(doc, "System in one sentence", "CFP Radar is a scheduled data-ingestion and quality-control system that converts fragmented academic opportunity sources into a current, evidence-backed catalog served through a Next.js product interface.")
    add_heading(doc, "Document map", 2)
    for item in [
        "Sections 1–3 explain the product and the five-layer system in plain language.",
        "Sections 4–8 explain discovery, integrity checks, deadlines, links, and deduplication.",
        "Sections 9–11 explain configuration, the technology stack, and the user-facing website.",
        "Sections 12–16 explain scale, tradeoffs, alternatives, risks, the roadmap, and market positioning.",
        "The final pages are an operator checklist, repository map, and architecture takeaway.",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()

    add_heading(doc, "1. The product idea", 1)
    add_body(doc, "Researchers face a fragmented market. Conference calls live on event websites, workshop calls live on separate pages, OpenReview contains submission portals, journals publish special issues, university presses announce book or chapter opportunities elsewhere, and reviewer recruitment may appear only in a conference update. Deadlines change, links break, and low-quality calls can look convincing.")
    add_body(doc, "CFP Radar turns that messy market into one trusted destination. Instead of repeatedly asking Google or an LLM what is open—and still risking a missed or stale call—a researcher can begin with one continuously monitored catalog. The promise is not merely “a big list.” It is “a current, trustworthy, navigable list with working official links.”")
    add_heading(doc, "What users can do", 2)
    for item in [
        "Browse open conferences, workshop paper calls, journals, and special issues.",
        "See a countdown based on the user’s own device timezone.",
        "Open the official CFP website and separately open the submission portal.",
        "Find conferences asking academics to propose and host workshops.",
        "Review calls for scholarly books and chapters from credible presses and publishers.",
        "Find conferences and workshops currently inviting academics or qualified students to review papers.",
        "Search with either keywords or a natural prompt such as “trustworthy AI with a deadline at least a month away.”",
        "Optionally share a manuscript with the AI assistant for venue fit, revision guidance, or ideas for extending existing work.",
        "Search and filter by opportunity type, domain, source, and legitimacy status.",
    ]:
        add_bullet(doc, item)
    add_callout(doc, "Product principle", "Correctness beats volume. A missing opportunity is inconvenient; a fake venue, wrong deadline, or misleading link can waste months of a researcher’s work.", LIGHT_BLUE)

    add_heading(doc, "The five tabs are five different jobs", 2)
    add_table(doc,
        ["Tab", "User job", "Why it is separate"],
        [
            ("CFP Dashboard", "Find somewhere to submit a paper.", "Paper calls have submission deadlines and venue-level trust checks."),
            ("Host a Workshop", "Find conferences inviting workshop proposals.", "This is for organizers, not authors submitting to an accepted workshop."),
            ("Books & Chapters", "Find credible editorial and press opportunities.", "Publisher legitimacy and proposal type differ from conference vetting."),
            ("Review Papers", "Find venues recruiting paper or ethics reviewers.", "Reviewer applications are separate from calls to submit or organize."),
            ("Paper → Venue", "Get manuscript-aware venue and revision support.", "This optional, user-controlled AI layer works only against the trusted opportunity catalog; discovery itself does not require manuscript access."),
        ],
        [1800, 3300, 4260]
    )
    add_heading(doc, "2. System architecture at a glance", 1)
    add_body(doc, "CFP Radar is organized into five boundaries. Each boundary has a distinct responsibility and failure mode.")
    add_step(doc, "Source adapters", "collect candidate opportunities from OpenReview, official conference sites, publishers, university presses, watchlists, and web leads.")
    add_step(doc, "Admission and policy layer", "evaluates currency, scope, legitimacy, source quality, deadline evidence, and opportunity type before publication.")
    add_step(doc, "Persistence layer", "stores accepted records in structured JSON and keeps a last-good OpenReview snapshot for recovery.")
    add_step(doc, "Serving layer", "reads the persisted state, enforces lifecycle rules, resolves links, deduplicates, sorts, and returns frontend-ready API data.")
    add_step(doc, "Presentation layer", "renders cards, filters, countdowns, tabs, and timezone-localized deadlines in React.")
    add_callout(doc, "Architectural implication", "User-facing requests do not initiate broad upstream discovery. Scheduled jobs absorb network latency and source instability; the serving path reads already-prepared state.")

    add_heading(doc, "The end-to-end flow", 2)
    add_body(doc, "External sources → scheduled discovery → integrity gates → deduplication and lifecycle checks → JSON stores → Next.js API → React dashboard → user clicks the official CFP or submission link.", keep=True)
    add_body(doc, "This separation also makes failures easier to contain. If OpenReview is temporarily unavailable, the website can still use the last successful snapshot. If an enrichment job fails, the basic verified call can remain visible rather than disappearing.")
    doc.add_page_break()

    add_heading(doc, "3. What happens every two days", 1)
    add_body(doc, "GitHub Actions is the scheduler. At 04:17 UTC every second day, it starts a workflow defined in .github/workflows/pipeline.yml. You can also trigger the same workflow manually.")
    add_step(doc, "Archive closed calls", "Remove any opportunity whose exact deadline has passed and retain a historical record in data/closed-calls.json.", 43)
    add_step(doc, "Synchronize OpenReview", "Read the homepage’s Open for Submissions directory, fetch exact invitation deadlines and venue metadata in batches, and save a last-good snapshot.", 43)
    add_step(doc, "Workshop paper discovery", "Check configured parent conferences for their individual workshops and challenges.", 43)
    add_step(doc, "Check the watchlist", "Look for newly opened calls associated with conferences and journals the team wants to monitor.", 43)
    add_step(doc, "Find calls to host workshops", "Search official parent-conference pages for organizer proposal deadlines and save them in a separate store.", 43)
    add_step(doc, "Search other conference systems", "Inspect CMT, EasyChair, HotCRP, custom CFP pages, and carefully controlled WikiCFP leads.", 43)
    add_step(doc, "Search books and chapters", "Check configured credible presses and publishers for relevant book and chapter proposal calls.", 43)
    add_step(doc, "Find calls for reviewers", "Check configured official conference and workshop pages for explicit current reviewer recruitment and a linked application route.", 43)
    add_step(doc, "Verify deadlines", "Re-read official pages and compare their dates with the stored values.", 43)
    add_step(doc, "Refresh legitimacy evidence", "Check history, indexing, domains, publishers, and red flags.", 43)
    add_step(doc, "Audit category coverage", "Fail the run unless papers, journals, workshop proposals, books, chapters, and reviewer calls all recorded a fresh check during this pipeline execution.", 43)
    add_step(doc, "Prune again", "Remove anything that closed while the workflow was running or was proven invalid during verification.", 43)
    add_step(doc, "Commit the refreshed catalog", "Save changed data files to Git. If the repository is connected to Vercel, that commit causes a fresh deployment.", 43)
    add_callout(doc, "Why prune twice", "The pipeline can take time, and verification can reveal that a call is already closed. A cleanup pass at both ends prevents stale records from surviving the run.", LIGHT_BLUE)
    doc.add_page_break()

    add_heading(doc, "4. How OpenReview works in this system", 1)
    add_body(doc, "OpenReview is valuable because it exposes a public list of venues that are open for submissions and an API containing exact invitation deadlines. But an OpenReview group page is usually a submission portal, not the full CFP.")
    add_heading(doc, "The two-phase strategy", 2)
    add_body(doc, "Phase 1 is the safety-critical fast path. It reads every homepage call, asks the API for exact deadlines, and fetches venue metadata such as the official website. Invitation requests and venue-metadata requests run concurrently in bounded batches. The result is saved immediately.")
    add_body(doc, "Phase 2 is best-effort enrichment. It visits official venue pages when needed, checks names and deadlines, and improves the evidence. If this slower phase fails, phase 1 remains available.")
    add_heading(doc, "Why the links are now separated", 2)
    for item in [
        "Main submission and OpenReview buttons point to the OpenReview group.",
        "View conference CFP or View workshop CFP prefers the verified official website embedded in OpenReview metadata.",
        "If no credible external CFP exists, the product should say the CFP link is unavailable rather than pretending the submission portal is the CFP.",
    ]:
        add_bullet(doc, item)
    add_body(doc, "For the Computational Psycholinguistics Meeting example, OpenReview stores https://cpl2026.sites.uu.nl/ as the venue website. The dashboard now uses that URL for View conference CFP while retaining OpenReview for submission.")
    add_heading(doc, "The last-good snapshot", 2)
    add_body(doc, "data/openreview-last-good.json is an insurance policy. The page API can reconstruct OpenReview cards from the last successful snapshot without making network calls during a user request. Expired entries are still removed by exact time, so a snapshot cannot keep a closed call alive forever.")
    add_callout(doc, "Design choice", "The live website trusts a recently prepared snapshot more than a fragile live scrape performed while a user is waiting. This is a standard reliability pattern: move expensive, failure-prone work out of the request path.")
    doc.add_page_break()

    add_heading(doc, "5. Other discovery sources", 1)
    add_heading(doc, "Watchlists", 2)
    add_body(doc, "data/watchlist.json is a private monitoring list, not a list of cards. Adding a conference name means “watch this,” not “publish this.” A card appears only after the pipeline finds a current official call with a real future deadline and the candidate passes admission checks.")
    add_heading(doc, "WikiCFP", 2)
    add_body(doc, "WikiCFP is treated as a lead generator, similar to a tip from a search engine. Its page alone is not proof. The pipeline must resolve the lead to an official event or publisher page, find the deadline there, confirm the venue’s identity and history, and then pass the integrity policy.")
    add_heading(doc, "CMT, EasyChair, HotCRP, PaperPlaza, and custom sites", 2)
    add_body(doc, "These platforms do not offer one uniform public directory like OpenReview. The system therefore combines configured watchlists, search, official-site resolution, and stricter evidence checks. Platform presence helps, but it does not automatically prove legitimacy.")
    add_heading(doc, "Workshops", 2)
    add_body(doc, "Two similarly named opportunities must never be mixed:")
    add_bullet(doc, "A workshop paper call asks authors to submit a paper to a workshop that already exists.")
    add_bullet(doc, "A workshop proposal call asks academics to propose and organize a new workshop for a parent conference.")
    add_body(doc, "They use separate source lists, scripts, APIs, stores, and tabs. The CFP importer also recognizes phrases such as Workshop Proposals and prevents those organizer calls from leaking into the paper dashboard.")
    add_heading(doc, "Books and chapters", 2)
    add_body(doc, "The book pipeline focuses on computer science, data science, AI/ML, and relevant interdisciplinary areas. It searches configured university presses and established scholarly publishers. A candidate needs a reachable official publisher or press link, evidence of a genuine proposal process, topical relevance, and a current deadline when the call is finite. Rejected candidates are retained separately so a human can audit what was dismissed.")
    add_heading(doc, "Calls for reviewers", 2)
    add_body(doc, "The Review Papers tab is for conferences and workshops inviting people to review submissions, including specialist roles such as ethics reviewers. The pipeline checks an allowlisted official venue page for explicit current language such as Call for Reviewers, seeking reviewers, or volunteer to review. It then confirms that the application route is linked from that page.")
    add_callout(doc, "Why the official page matters", "A Google Form, Microsoft Form, or shortened link is only the place where the user applies. Anyone can create a form, so the form is never accepted as proof that the opportunity is legitimate. The official conference or workshop page provides that proof.", LIGHT_BLUE)

    add_heading(doc, "6. The integrity and admission gates", 1)
    add_body(doc, "The system uses several modest checks together rather than trusting one magic score. This is closer to a hiring process than a spam filter: identity, evidence, history, and current status all matter.")
    add_heading(doc, "A candidate should answer yes to these questions", 2)
    for item in [
        "Is the opportunity in the product’s subject scope?",
        "Is there an official or independently credible source page?",
        "Does that page show a real future submission or proposal deadline, or explicitly state a legitimate rolling process?",
        "Does the venue or publisher identity match across sources?",
        "Is the link safe and reachable?",
        "For reviewer calls, does the official venue page explicitly recruit reviewers and link the application route?",
        "Is there credible history, indexing, society, university, or publisher evidence?",
        "Are there no major red flags such as name imitation, strange domains, unsupported sponsor claims, or predatory-publisher patterns?",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Signals used", 2)
    add_body(doc, "CORE and SCImago-style classifications, established publisher or society domains, Crossref proceedings history, OpenAlex indexing evidence, official university ownership, and OpenReview’s current submission invitation can all contribute. No single signal is perfect, so the system combines them.")
    add_heading(doc, "What the optional LLM does—and does not do", 2)
    add_body(doc, "The opportunity catalog does not need a user’s manuscript. Separately, a researcher may opt in to an AI-assisted workflow that compares a draft with verified open venues, identifies possible revision priorities, or helps think through how existing work might be extended. A hosted language model can also help interpret messy source text, but it is not the source of truth. Deterministic rules, official links, exact timestamps, and structured source evidence remain the final control points.")
    add_body(doc, "Manuscript access must be explicit rather than automatic. The product should say what text will leave the application, which model provider will receive it, and whether it is retained. Users should still be able to use the entire discovery dashboard without exposing a draft to any LLM.")
    add_callout(doc, "Product safety rule", "AI can help interpret evidence; it should not invent the evidence. If the official source cannot support a claim, the product should show uncertainty or reject the candidate.")

    add_heading(doc, "7. Deadlines, timezones, and automatic removal", 1)
    add_body(doc, "Every finite call is stored as a full timestamp, not merely “July 18.” A timestamp includes the date, time, and timezone offset. Internally it can be compared as one universal instant.")
    add_heading(doc, "What the user sees", 2)
    add_body(doc, "The frontend converts the stored instant into the timezone reported by the user’s device. A person in New York sees Eastern time; someone in California sees Pacific time. Both are looking at the same deadline instant.")
    add_heading(doc, "One deadline rule, enforced in three places", 2)
    add_body(doc, "Assume a conference deadline is 8:00 PM Eastern. The card should become unavailable at 8:00 PM; it should not remain visible until the next two-day pipeline run. The system applies that same deadline timestamp in three contexts:")
    add_bullet(doc, "Already-open dashboard: the shared browser clock causes the card list to re-evaluate, so the card disappears without the user manually refreshing the page.")
    add_bullet(doc, "New page load or refresh: /api/cfps compares the deadline with the current server time and does not return the expired card.")
    add_bullet(doc, "Scheduled data cleanup: the next pipeline run permanently removes the expired record from data/cfps.json and moves it into the closed-calls archive.")
    add_body(doc, "Therefore, at 8:02 PM the card can already be absent from the product even though the JSON file has not yet been cleaned. The browser and API control what users can currently see; the scheduled pipeline keeps the persisted catalog tidy.")
    add_callout(doc, "Why all three checks exist", "They cover three different states—an existing browser session, a fresh request, and stored data—but they use the same deadline and the same open-versus-closed rule.", LIGHT_BLUE)
    add_heading(doc, "Rolling and unknown deadlines", 2)
    add_body(doc, "A rolling journal can stay visible when the publisher explicitly accepts submissions continuously. A finite call without a real deadline is normally rejected. A special OpenReview status may say that submissions are active but the public deadline is unavailable; the UI should label that uncertainty instead of inventing a countdown.")
    add_body(doc, "Reviewer recruitment often has no published application deadline. In that case CFP Radar does not invent one. Instead, the card receives a short verification lease. The two-day pipeline renews that lease only while the official venue page still explicitly recruits reviewers. A successfully checked page that removes or closes the call removes the card, while repeated failed checks cannot preserve it indefinitely.")
    add_callout(doc, "Common trap", "A date without a timezone is ambiguous. The system should preserve the timezone stated by the official source or use a clearly documented convention such as Anywhere on Earth only when the source says so.", LIGHT_BLUE)

    add_heading(doc, "8. Deduplication and grouping", 1)
    add_body(doc, "The same call can arrive through OpenReview, an official site, a watchlist, and a search result. Without deduplication, users see repeated cards. But overly aggressive deduplication is also dangerous because two unrelated venues can share an acronym.")
    add_heading(doc, "The current identity strategy", 2)
    for item in [
        "Exact OpenReview group IDs are the strongest identity for OpenReview records.",
        "Canonicalized titles and years help match a live OpenReview call to an existing curated card.",
        "Fuzzy acronym matching is not allowed between two live OpenReview cards.",
        "Tracks belonging to one venue can be grouped into one card with multiple real deadlines.",
        "A track countdown is shown only when it differs from the card’s primary deadline, preventing duplicate countdowns.",
        "Deduplication keys include opportunity type and deadline so different calls are not casually collapsed.",
    ]:
        add_bullet(doc, item)
    add_body(doc, "The practical lesson is that identity is a product concept, not merely a string comparison. “Same acronym” does not mean “same opportunity.”")
    add_heading(doc, "Where this can still improve", 2)
    add_body(doc, "A future database could assign stable venue IDs, event-edition IDs, and call IDs. That would make it easier to represent one conference with several tracks, extensions, and submission rounds without relying on names. The current JSON model is sufficient for the present scale but less expressive than a relational model.")
    doc.add_page_break()

    add_heading(doc, "9. Configuration versus core logic", 1)
    add_body(doc, "Your design goal is correct: adding a professor’s watchlist, a new conference family, a publisher, or a ranking label should usually change settings—not the engine.")
    add_table(doc,
        ["File", "What a product operator changes", "What should happen"],
        [
            ("data/watchlist.json", "Conference and journal names to monitor", "Discovery watches them; nothing is published without an open verified call."),
            ("data/workshop-sources.json", "Parent conferences whose workshop paper calls matter", "The workshop discovery job checks their official directories."),
            ("data/workshop-proposal-sources.json", "Parents that may invite workshop organizers", "Results feed the Host a Workshop tab."),
            ("data/book-call-sources.json", "Trusted presses and publisher sources", "The book pipeline checks configured sources and applies its admission policy."),
            ("data/reviewer-call-sources.json", "Official venues that may recruit reviewers", "Results feed the Review Papers tab only after the official-page and application-link checks."),
            ("data/catalog-config.json", "Labels, venue types, ranking labels, and link-policy patterns", "Frontend behavior changes without editing React components."),
            ("Environment variables", "Concurrency, timeouts, models, and API keys", "Operations can tune performance and services without code changes."),
        ],
        [2500, 3400, 3460]
    )
    add_heading(doc, "What is still core logic", 2)
    add_body(doc, "Rules such as “an expired finite call is never active,” “WikiCFP is not authoritative,” “OpenReview IDs are exact identities,” “organizer proposals do not belong in the paper tab,” and “an external reviewer form is not legitimacy evidence” are product invariants. Those belong in code and tests because every lab should inherit the same safety behavior.")
    add_heading(doc, "The next configuration upgrade", 2)
    add_body(doc, "JSON configuration is developer-friendly but not yet product-friendly. A future admin screen could let a lab owner add sources, domains, and watchlists through forms. The backend would validate those settings against a schema and store them per organization.")
    doc.add_page_break()

    add_heading(doc, "10. Technology stack and why it was chosen", 1)
    add_table(doc,
        ["Technology", "Role", "Why it fits now", "Tradeoff / alternative"],
        [
            ("Next.js 14", "Web application and API routes", "Frontend and backend live in one JavaScript project; easy Vercel deployment.", "A separate FastAPI or Django backend gives stronger service boundaries but adds deployment complexity."),
            ("React", "Interactive tabs, filters, cards, and countdowns", "Component model makes repeated UI predictable and reusable.", "A simpler server-rendered site would use less client JavaScript but feel less interactive."),
            ("Plain CSS", "Visual design and responsive layout", "Small dependency surface and easy theme variables.", "Tailwind or a component library can speed consistent design at the cost of another abstraction."),
            ("Node.js scripts", "Discovery, verification, merging, and cleanup", "Same language as the web app; excellent async networking.", "Python has a richer data ecosystem, but a second language raises maintenance cost."),
            ("JSON files", "Catalogs, configuration, snapshots, and rejection logs", "Transparent, version-controlled, inexpensive, and enough for a shared catalog.", "PostgreSQL is better for accounts, per-lab settings, audit queries, and concurrent writes."),
            ("GitHub Actions", "Two-day scheduler and automation runner", "No always-on worker server; workflows and logs live beside the code.", "A job queue such as Temporal, Celery, or cloud queues offers retries and observability at higher complexity."),
            ("Vercel-compatible deployment", "Hosts the Next.js app", "Fast setup and automatic deploys from Git commits.", "A container platform gives more runtime control and is better for long background jobs."),
            ("OpenReview API", "Current calls, invitation deadlines, and official venue metadata", "Structured primary data and stable group identities.", "It covers only venues using OpenReview and can change behavior or rate limits."),
            ("Crossref and OpenAlex", "History and indexing evidence", "Open scholarly metadata reduces manual checks.", "Coverage is imperfect; paid indexing sources may be more authoritative."),
            ("Groq or NVIDIA LLM, optional", "Messy-text extraction and paper recommendations", "Hosted models avoid operating your own GPU.", "Quality, cost, and privacy depend on the provider; deterministic parsing should remain the backbone."),
            ("pdf-parse", "Extracts text from uploaded papers", "Lightweight local PDF text extraction in Node.", "Complex scanned or multi-column PDFs may need OCR or a specialist parser."),
        ],
        [1700, 2100, 2850, 2710]
    )
    add_callout(doc, "Why JavaScript everywhere", "For a small team, one language across the browser, API, and scheduled scripts lowers the number of concepts you must maintain. It is an organizational choice as much as a technical one.")

    add_heading(doc, "11. What happens when a user opens the site", 1)
    add_step(doc, "Next.js serves the page shell", "The user receives the tabs, layout, styles, and React application.", 44)
    add_step(doc, "The dashboard calls /api/cfps", "This is a quick local data read, not a live crawl of the internet.", 44)
    add_step(doc, "The backend reconstructs safe current data", "It reads data/cfps.json, optionally merges the last-good OpenReview snapshot, removes expired records, deduplicates, resolves official CFP links, and sorts by deadline.", 44)
    add_step(doc, "React renders cards", "The browser applies search, filters, sorting, and local-time formatting.", 44)
    add_step(doc, "One shared clock updates lifecycle state", "Cards do not each create an independent timer. This reduces duplicate work and makes expiry behavior consistent.", 44)
    add_step(doc, "A button opens the correct destination", "CFP buttons prefer official external pages; OpenReview buttons remain submission links.", 44)
    add_heading(doc, "Why the API does not live-fetch OpenReview", 2)
    add_body(doc, "Doing so would make every page load depend on hundreds of remote calls, expose users to rate limits, and create inconsistent counts. A scheduled snapshot is more predictable, easier to audit, and much cheaper.")
    add_heading(doc, "Caching choice", 2)
    add_body(doc, "The CFP API currently sends no-store, meaning the browser should not rely on an old API response. The expensive upstream work is already cached as data files. A future production deployment could safely add a short CDN cache if invalidation at deadline boundaries is handled carefully.")
    add_heading(doc, "Hybrid search without sharing a manuscript", 2)
    add_body(doc, "A researcher may want better discovery without uploading a paper. The dashboard search therefore accepts both ordinary keywords—such as KDD or explainable AI—and prompt-like requests. For example: “My work is on trustworthy AI and I need a deadline at least a month away” or “Deadline in less than 30 days.” The query stays in the browser and does not call an LLM.")
    add_heading(doc, "What happens to the query", 3)
    add_bullet(doc, "Intent parsing extracts hard constraints such as opportunity type, trusted-only status, rolling calls, and relative or named deadline windows.")
    add_bullet(doc, "Semantic expansion connects a limited set of research concepts. For example, trustworthy AI can include fairness, accountability, explainability, robustness, privacy, and governance.")
    add_bullet(doc, "BM25 ranks the current cards. Rare and informative terms receive more weight, repetition has diminishing value, and long descriptions are normalized so they do not dominate focused records.")
    add_bullet(doc, "Exact phrases and important fields—name, acronym, topics, domain, conference, and publisher—receive additional weight. Existing tab and dropdown filters still apply.")
    add_body(doc, "The BM25 index is calculated from whatever cards are currently loaded. When the pipeline adds a conference, special issue, workshop proposal, book call, or reviewer opportunity, that record becomes searchable automatically; no per-source search configuration is required.")
    add_callout(doc, "How this helps", "The search box becomes a lightweight decision interface. Users can express both subject fit and practical constraints without exposing unpublished work: “responsible AI journals,” “workshops within three weeks,” “chapter proposals about AI,” or “ethics reviewer opportunities.”", LIGHT_BLUE)
    add_heading(doc, "Honest assessment: was BM25 overengineering?", 2)
    add_body(doc, "For the earliest version of the product, yes—BM25 would have been unnecessary. When the catalog was small and users were typing acronyms, a case-insensitive contains check was easier to understand and entirely adequate. Building a ranking system before observing prompt-style search behavior would have added complexity without proven value.")
    add_body(doc, "For the current feature, BM25 is a proportionate choice rather than serious overengineering. The product now searches hundreds of records with uneven text length, and users expect the most relevant cards first. BM25 adds better ranking without a hosted model, external search service, vector database, API key, or privacy exposure. Its implementation and test burden are still modest.")
    add_body(doc, "TF-IDF plus cosine similarity would also have been reasonable. BM25 fits short queries and variable-length cards slightly better because it explicitly controls term saturation and document-length normalization. Neither algorithm understands dates or research intent by itself, which is why constraint parsing and semantic expansion remain separate parts of the hybrid design.")
    add_callout(doc, "Where overengineering would begin", "Adding embeddings, a vector database, or an LLM call on every keystroke is not justified yet. Those options should be considered only after real search logs show that BM25 plus the intent parser repeatedly misses conceptual queries that matter to users.", LIGHT_BLUE)
    doc.add_page_break()

    add_heading(doc, "12. Can this support 100–200 users?", 1)
    add_body(doc, "Yes, if “users” means 100–200 people reading the same shared catalog. The page request is mostly a JSON read and filter. That traffic is modest for Next.js and Vercel-class hosting.")
    add_heading(doc, "What scales well already", 2)
    for item in [
        "Discovery work runs once per schedule instead of once per user.",
        "Network calls use bounded asynchronous pools, so the pipeline is faster without flooding source sites.",
        "Static JSON is cheap to read and easy to distribute.",
        "Last-good snapshots prevent one external outage from becoming a frontend outage.",
        "Config files let several maintainers share one engine.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "What “100–200 labs” would change", 2)
    add_body(doc, "If each professor or lab needs a private watchlist, different domains, notifications, saved papers, or admin permissions, this becomes a multi-tenant SaaS product. JSON files in one repository are no longer the right source of truth.")
    add_heading(doc, "A sensible multi-tenant architecture", 2)
    for item in [
        "User accounts and organization workspaces.",
        "PostgreSQL tables for venues, calls, sources, organizations, watchlists, and subscriptions.",
        "A background job queue for discovery, verification, retries, and per-lab notifications.",
        "An admin UI instead of asking users to edit JSON.",
        "Role-based access such as owner, editor, and viewer.",
        "Email or Slack digests built from each lab’s saved interests.",
        "Central observability for failed sources, stale data, and pipeline duration.",
    ]:
        add_bullet(doc, item)
    add_callout(doc, "Do not overbuild too early", "A shared Penn State pilot should validate which customization and notification features people actually use. Move to the database and account model when separate ownership becomes a real requirement, not merely a hypothetical one.", LIGHT_BLUE)
    doc.add_page_break()

    add_heading(doc, "13. Alternatives and tradeoffs", 1)
    add_heading(doc, "Alternative A: live crawl on every page load", 2)
    add_body(doc, "This sounds fresh but is a poor fit. It is slow, rate-limit prone, and impossible to make consistent across users. The scheduled snapshot is the better product choice.")
    add_heading(doc, "Alternative B: database from day one", 2)
    add_body(doc, "PostgreSQL would improve relationships, querying, accounts, and audit history. The cost is migrations, hosting, backups, permissions, and more operational concepts. JSON is a good prototype and shared-catalog choice; a database becomes better once personalization or concurrent editing arrives.")
    add_heading(doc, "Alternative C: Python data pipeline plus JavaScript frontend", 2)
    add_body(doc, "Python could provide stronger scraping, ML, and data tooling. The tradeoff is maintaining two ecosystems. This is attractive later if discovery becomes a substantial data platform with OCR, browser automation, or statistical ranking.")
    add_heading(doc, "Alternative D: commercial search and indexing APIs", 2)
    add_body(doc, "Paid services can improve search quality and bibliographic authority. They raise operating cost and create vendor dependence. A reasonable path is to start with open sources, measure false positives and missed calls, then pay for the specific evidence layer that has the highest value.")
    add_heading(doc, "Alternative E: browser automation for every source", 2)
    add_body(doc, "A headless browser can read JavaScript-heavy pages, but it is slower and more fragile than an API or normal HTTP fetch. The present system correctly prefers structured APIs and official page text, using heavier techniques only where necessary.")
    add_heading(doc, "Alternative F: let an LLM run the whole workflow", 2)
    add_body(doc, "An agent could search broadly and reason flexibly, but outputs would be harder to reproduce, test, and audit. The better architecture uses deterministic rules for identity, time, safety, and lifecycle, with AI as an assistant around messy language.")
    doc.add_page_break()

    add_heading(doc, "14. Known limitations and operational risks", 1)
    for title, detail in [
        ("Source changes", "OpenReview or publisher page formats can change. Parser tests and last-good snapshots reduce damage, but maintainers must monitor failures."),
        ("Official pages can be ambiguous", "A page may mix abstract, paper, notification, camera-ready, and workshop dates. The verifier must label the deadline type rather than choosing the first date."),
        ("Coverage is not exhaustive", "Some calls are announced only in mailing lists, newsletters, social media, or PDFs that search engines have not indexed."),
        ("Reviewer calls may omit deadlines", "The verification lease prevents a previously valid reviewer form from remaining visible forever when the official venue stops recruiting."),
        ("Open metadata is incomplete", "Crossref and OpenAlex do not perfectly represent every legitimate venue, especially new or regional events."),
        ("JSON has weak concurrency", "Two automated jobs or human editors can create merge conflicts. One coordinated pipeline avoids most of this at the current scale."),
        ("No per-user preferences yet", "Everyone sees the same catalog and filters locally. Saved watchlists and notifications need accounts and durable storage."),
        ("Uploaded-paper privacy", "If the recommender sends extracted paper text to a hosted LLM, users need a clear privacy notice and should avoid confidential manuscripts unless the provider terms are acceptable."),
        ("Email notifications are absent", "The dashboard is pull-based today. Researchers must visit it unless a digest feature is added."),
        ("Prompt search is bounded", "The local intent parser supports common deadline and opportunity patterns, not arbitrary natural-language reasoning. Search failures should be measured before adding embeddings or an LLM dependency."),
    ]:
        add_body(doc, f"{title}: {detail}", bold_lead=title + ":")
    add_callout(doc, "Best monitoring metric", "Track precision and freshness, not just card count: percentage of cards with working official CFP links, percentage with verified deadlines, duplicates found, expired cards served, and rejected candidates later approved by a human.")
    doc.add_page_break()

    add_heading(doc, "15. Recommended product roadmap", 1)
    add_heading(doc, "Phase 1 — make the shared catalog trustworthy", 2)
    for item in [
        "Add a visible “last checked” timestamp and evidence detail for every card.",
        "Create an operator review queue for uncertain or rejected candidates.",
        "Measure broken-link rate, deadline mismatches, duplicate rate, and source coverage.",
        "Add automated alerts when a source parser suddenly returns far fewer records.",
        "Finish schema validation for every configuration and output file.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Phase 2 — learn what labs personalize", 2)
    for item in [
        "Allow saved filters and a lightweight personal watchlist.",
        "Offer a weekly email digest before building complex notification rules.",
        "Interview several professors and graduate students about how they triage opportunities.",
        "Track which links and tabs users actually open.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Phase 3 — become multi-tenant only if demand proves it", 2)
    for item in [
        "Introduce accounts, lab workspaces, PostgreSQL, permissions, and organization settings.",
        "Move long-running discovery to a queue with durable retries and dashboards.",
        "Provide an admin interface for sources, watchlists, domains, and trust policy.",
        "Add Slack, Teams, or email integrations after notification preferences are understood.",
    ]:
        add_bullet(doc, item)
    add_callout(doc, "Product recommendation", "Do not judge success by “we found 5,000 calls.” Judge it by “a researcher trusted the system enough to make a submission decision, and the source, deadline, and opportunity type were correct.”", LIGHT_BLUE)
    doc.add_page_break()

    add_heading(doc, "16. Product positioning: why CFP Radar should exist", 1)
    add_heading(doc, "The USP in one sentence", 2)
    add_callout(doc, "Positioning statement", "CFP Radar is the trusted one-stop destination for currently open academic opportunities across conferences, journals, special issues, workshops, books, chapters, and peer-review service—so researchers do not have to repeatedly search Google, ask an LLM, or depend on scattered announcements.", LIGHT_BLUE)
    add_body(doc, "The product is not trying to replace or compete with OpenReview. OpenReview is excellent submission and peer-review infrastructure and is also a valuable structured source. CFP Radar solves the problem before submission: finding the right legitimate opportunity in time, regardless of where its official call or submission portal happens to live.")
    add_body(doc, "A user should be able to answer four questions quickly: What credible opportunities are open? Which ones are relevant to my work? When do they close in my timezone? Where is the correct official page to read or act?")
    add_heading(doc, "What is actually differentiated", 2)
    for item in [
        "One destination instead of repeated searching. The system continuously monitors fragmented sources so a researcher does not have to reconstruct the market from Google results, newsletters, bookmarks, and one-off LLM answers.",
        "Actionability over listing volume. Closed finite calls disappear at their exact deadline, and a card distinguishes the official CFP from the submission portal.",
        "Evidence-backed trust. Aggregators such as WikiCFP can generate useful leads, but their listings are not accepted as proof. Publication requires an official source, matching identity, a verified deadline, and credible venue or publisher history.",
        "Several academic jobs in one focused product. Papers, workshop papers, calls to host workshops, journal special issues, books, chapters, and calls to review papers are separated by intent but available in one research-opportunity system.",
        "A CS/AI-specific quality bar. The catalog can prioritize the fields, rankings, societies, presses, and interdisciplinary edges that matter to the target labs instead of treating every event category equally.",
        "Shared infrastructure without hard-coded lab logic. Sources, watchlists, and labels are configuration; lifecycle, identity, and safety rules remain common product invariants.",
        "An optional path from discovery to manuscript decisions. If the user explicitly opts in, AI can connect a draft to verified opportunities, suggest revisions for a target venue, or help explore responsible extensions of existing work.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()
    add_heading(doc, "Who gets the clearest value", 2)
    add_table(doc,
        ["User", "Current pain", "Reason to use CFP Radar"],
        [
            ("Graduate researcher", "Repeatedly searches Google or asks an LLM, yet can still miss a call or receive stale information.", "Starts from one monitored catalog of open, source-backed opportunities and the correct next click."),
            ("Professor or PI", "Opportunities are passed around informally and become stale between lab meetings.", "Gives the lab one shared, current view with explainable trust signals."),
            ("Lab or program manager", "Maintains spreadsheets, mailing-list notes, and reminders by hand.", "Automates monitoring and expiry while keeping sources and watchlists configurable."),
            ("Workshop organizer", "Calls to propose workshops are mixed with calls to submit papers to existing workshops.", "Uses a dedicated organizer view with the correct proposal deadline and parent conference."),
            ("Interdisciplinary scholar", "Book and chapter calls are scattered across press pages and publisher announcements.", "Finds relevant scholarly calls inside the same evidence and lifecycle model."),
            ("Emerging reviewer", "Reviewer recruitment is easy to miss and an isolated form can be difficult to trust.", "Uses the official venue evidence, role details, and verified application route to find service opportunities."),
        ],
        [1900, 3600, 3860]
    )
    doc.add_page_break()

    add_heading(doc, "How CFP Radar fits into the existing ecosystem", 1)
    add_body(doc, "This is primarily an ecosystem and workflow comparison, not a competitor chart. CFP Radar should consume useful signals from existing platforms and send users to the correct platform when it is time to read the official call or submit.")
    add_table(doc,
        ["Tool or workflow", "Its role", "CFP Radar’s relationship to it"],
        [
            ("OpenReview", "Runs submission, peer review, discussion, and publishing through a structured platform and API.", "A complementary source and destination—not a competitor. CFP Radar discovers an OpenReview call, verifies and contextualizes it, then sends the researcher to OpenReview when appropriate."),
            ("WikiCFP", "Provides a broad community-maintained technology CFP directory and useful discovery leads.", "An untrusted lead source. Because listings may be stale, incorrect, duplicated, or low-confidence, CFP Radar publishes them only after confirmation from an official source and legitimacy checks."),
            ("Conference Alerts", "Provides broad event discovery and interest-based alerts; organizers can add and promote events.", "Another potential lead and notification channel. CFP Radar focuses on verified academic opportunities, correct action links, and exact open-call lifecycle."),
            ("Google, LLMs, newsletters, and mailing lists", "Surface new or niche announcements, but require repeated searching and manual verification.", "The workflow CFP Radar is meant to replace for routine discovery. These channels still provide leads, while CFP Radar turns them into a monitored and trusted catalog."),
            ("Lab spreadsheets or bookmarks", "Capture local priorities and informal knowledge.", "CFP Radar preserves local configuration while automating source checks, deduplication, deadlines, and shared presentation."),
        ],
        [1800, 3600, 3960]
    )
    add_body(doc, "Ecosystem sources checked in July 2026: OpenReview About, WikiCFP, and Conference Alerts. This describes public product scope only—not internal performance or business data.")
    add_heading(doc, "The honest product tradeoff", 2)
    add_body(doc, "The one-stop promise creates a real precision-versus-coverage problem. If admission is too loose, the catalog inherits the bogus and stale entries users are trying to escape. If it is too strict, legitimate niche calls may be missed. The product must expose coverage gaps, keep a human review path, and improve recall without weakening the trust threshold.")
    add_body(doc, "The core metric is therefore not total cards. It is the percentage of opportunities that are relevant, legitimate, open, correctly dated, and linked to the right destination—alongside the rate of important calls the system missed. Time saved per researcher or lab is another useful metric.")
    add_heading(doc, "Why a user would choose it", 2)
    add_body(doc, "A user chooses CFP Radar because opportunity discovery should not be a recurring research project. They want one dependable place to see what is open across the academic ecosystem, with enough evidence to act. Consistent expiry, integrity screening, opportunity-type separation, and official links reduce both search time and the chance of missing or trusting the wrong call.")
    add_callout(doc, "The adoption test", "If a researcher begins with CFP Radar instead of asking Google or an LLM to reconstruct the open-call landscape—and can act without re-checking every card from scratch—the product is delivering its USP.", LIGHT_BLUE)

    add_heading(doc, "Optional manuscript intelligence", 1)
    add_body(doc, "The trusted catalog is the core product. Manuscript intelligence is a separate opt-in layer for users who are comfortable exposing a draft to an LLM. It should support three practical decisions:")
    add_bullet(doc, "Venue fit: rank only verified, currently open opportunities and explain the topical, methodological, and contribution fit.")
    add_bullet(doc, "Revision planning: identify where a draft may need stronger framing, evidence, experiments, or positioning for a chosen venue—without promising acceptance.")
    add_bullet(doc, "Extension planning: help a researcher distinguish the new contribution from existing work and explore meaningful follow-up questions, datasets, analyses, or experiments—without inventing novelty.")
    add_callout(doc, "Privacy boundary", "The dashboard must remain fully useful without an upload. Manuscript sharing requires explicit consent, a clear provider and retention notice, and an understandable description of what text leaves the application.", LIGHT_BLUE)

    add_heading(doc, "What could become defensible", 1)
    add_body(doc, "Scraping pages and calling an LLM are not durable advantages by themselves. Other teams can reproduce those mechanics. The more defensible asset is the accumulated operating knowledge around academic opportunities:")
    for item in [
        "Reliable source adapters and tests for difficult conference, publisher, and press websites.",
        "A stable identity graph connecting a venue series, annual edition, track, workshop, proposal call, and submission portal.",
        "An evidence history showing which deadline and link were observed, when they changed, and why a call was admitted or rejected.",
        "A rejection corpus, human-review feedback, and lab-specific usage signals that improve precision and relevance without hiding decisions inside an opaque score.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "A focused go-to-market wedge", 2)
    add_body(doc, "The strongest initial position is not “all academic opportunities for everyone.” It is “the trusted one-stop opportunity destination for a small number of CS/AI labs.” Start with Penn State teams, measure whether the catalog prevents missed calls and changes real decisions, and expand disciplines or institutions only after the quality model works in that wedge.")
    add_heading(doc, "Questions the product team should keep testing", 2)
    for item in [
        "Will researchers trust a verified subset enough to prefer it over a larger unverified directory?",
        "Which missed calls are unacceptable, and which low-confidence calls should remain outside the catalog?",
        "Do users value immediate alerts, weekly digests, manuscript guidance, or shared lab curation most?",
        "What evidence must be visible on a card before a user stops re-checking the source from scratch?",
        "Can the system maintain its precision as sources, disciplines, and labs increase?",
    ]:
        add_bullet(doc, item)
    add_callout(doc, "Strategic summary", "CFP Radar should be the trusted one-stop discovery and decision destination above fragmented academic sources. OpenReview and other submission systems remain partners and destinations; optional manuscript AI deepens the decision support only when the user asks for it.")
    doc.add_page_break()

    add_heading(doc, "How to run CFP Radar", 1)
    add_body(doc, "This section is the practical setup path for a new machine. You need Node.js 20.x, npm, and a copy of the repository. The product catalog and hybrid search work without a paid service or API key.")
    add_heading(doc, "First local run", 2)
    add_code_block(doc, """npm install
cp .env.example .env.local
npm run dev""")
    add_body(doc, "Then open http://localhost:3000. Keep the terminal running while you use the local site. Stop it with Control+C.")
    add_callout(doc, "What the copy command does", "It creates a private local settings file from the safe example. Do not place real API keys in .env.example, and do not commit .env.local to Git.", LIGHT_BLUE)
    add_heading(doc, "Production-style check", 2)
    add_body(doc, "Before deploying or sharing a change, confirm that the optimized application can compile and start:")
    add_code_block(doc, """npm run build
npm run start""")
    add_body(doc, "Open http://localhost:3000 again. The build command catches problems that the development server may tolerate.")
    add_heading(doc, "Run the data pipeline locally", 2)
    add_body(doc, "The website reads the saved catalog immediately. The following commands refresh each source family and then apply verification, trust, and expiry rules. Run only the source you are changing during development; the scheduled GitHub workflow runs the complete sequence every two days.")
    add_code_block(doc, """npm run prune:closed
npm run scrape
npm run sync:workshops
npm run watchlist
npm run sync:workshop-proposals
npm run discover:web
npm run sync:book-calls
npm run sync:reviewer-calls
npm run verify
npm run legitimacy
npm run prune:closed""")
    add_body(doc, "These commands make network requests and can take several minutes. The first and last cleanup steps ensure that calls which expire before or during the run do not remain in the saved catalog.")
    add_heading(doc, "Tests to run before publishing", 2)
    add_code_block(doc, """npm run test:discovery
npm run test:openreview
npm run test:links
npm run test:workshop-proposals
npm run test:book-calls
npm run test:reviewer-calls
npm run test:lifecycle
npm run test:admission
npm run test:dedupe
npm run test:search
npm run build""")

    add_heading(doc, "What to put in .env.local", 1)
    add_body(doc, "Treat .env.example as documentation and a copyable template. Put machine-specific values and secrets in .env.local. For the normal dashboard, CFP cards, deadlines, links, tabs, and hybrid search, no value is required: the blank API-key fields can remain blank.")
    add_heading(doc, "Option A: dashboard only", 2)
    add_body(doc, "Copy the example file and change nothing. This is enough to browse calls, filter cards, use prompt-like search, and test the frontend.")
    add_heading(doc, "Option B: enable Groq manuscript features", 2)
    add_body(doc, "Create a Groq key, then put the real value only in .env.local:")
    add_code_block(doc, """LLM_PROVIDER=groq
GROQ_API_KEY=your_real_groq_key
LLM_MODEL=qwen/qwen3.8-27b
LLM_TIMEOUT_MS=180000""")
    add_heading(doc, "Option C: enable NVIDIA models", 2)
    add_body(doc, "Use this instead when you want a model exposed through NVIDIA NIM:")
    add_code_block(doc, """LLM_PROVIDER=nvidia
NVIDIA_API_KEY=your_real_nvidia_key
LLM_MODEL=meta/llama-3.1-8b-instruct
LLM_TIMEOUT_MS=180000""")
    add_callout(doc, "Choose one provider", "You do not need both keys. Groq is the simplest default. NVIDIA is useful only when you intentionally select one of the NVIDIA-hosted model options.", LIGHT_BLUE)
    doc.add_page_break()
    add_heading(doc, "Which settings normally need editing", 2)
    add_table(doc,
        ["Variable group", "When you need it", "Normal action"],
        [
            ("LLM_PROVIDER and one API key", "Only for manuscript recommendation, draft review, extension ideas, and the optional LLM trust opinion.", "Choose Groq or NVIDIA and add one real key to .env.local."),
            ("LLM_MODEL and LLM_TIMEOUT_MS", "When selecting a different supported model or allowing longer responses.", "Keep the example defaults unless you have a reason to change them."),
            ("OPENREVIEW_*", "When tuning live OpenReview concurrency, retries, cache duration, or timeouts.", "Keep the defaults for normal development."),
            ("WATCHLIST_*, DISCOVERY_*", "When tuning how many external candidates are checked and how many requests run together.", "Keep the defaults; reduce concurrency if a source rate-limits your machine."),
            ("BOOK_DISCOVERY_*", "When tuning book and chapter discovery breadth and concurrency.", "Keep the defaults; source names and domains belong in data/book-call-sources.json."),
            ("REVIEWER_DISCOVERY_CONCURRENCY", "When tuning how many official reviewer-call pages are checked in parallel.", "Keep the default; monitored venues belong in data/reviewer-call-sources.json."),
        ],
        [2500, 3800, 3060]
    )
    add_heading(doc, "Environment values and command-line scripts", 2)
    add_body(doc, "Next.js automatically reads .env.local when you run the website. Plain Node.js pipeline scripts do not automatically read that file. If a local script needs your optional LLM key, load the values into that terminal first:")
    add_code_block(doc, """set -a
source .env.local
set +a
npm run verify""")
    add_body(doc, "GitHub Actions receives the same optional values from repository variables and secrets. Put GROQ_API_KEY or NVIDIA_API_KEY in GitHub Secrets, not in the repository. The structured source, deadline, ranking, and publisher checks still run without an LLM key.")
    add_callout(doc, "No account setup required", "The current product has no login, email verification, Supabase, Resend, or WhatsApp configuration. Do not add those environment variables unless that feature is intentionally reintroduced later.", LIGHT_BLUE)
    doc.add_page_break()

    add_heading(doc, "Operator’s guide: safe changes without editing core logic", 1)
    add_heading(doc, "To monitor another conference or journal", 2)
    add_step(doc, "Edit data/watchlist.json", "Add the name and relevant identifiers using the existing record pattern.", 45)
    add_step(doc, "Run the watchlist or full pipeline", "The watchlist is monitoring configuration, not automatic publication approval.", 45)
    add_step(doc, "Review rejections and output", "Confirm the candidate found an official page and passed the deadline and legitimacy gates.", 45)
    add_heading(doc, "To monitor workshop paper calls", 2)
    add_body(doc, "Add the parent conference to data/workshop-sources.json. Do not add its call for workshop organizers as a workshop paper CFP.")
    add_heading(doc, "To monitor calls for workshop organizers", 2)
    add_body(doc, "Use data/workshop-proposal-sources.json. Those records belong to the Host a Workshop tab and have their own lifecycle.")
    add_heading(doc, "To add book and chapter sources", 2)
    add_body(doc, "Use data/book-call-sources.json and preserve the trusted-publisher metadata pattern. The discovery engine should still verify each actual call rather than trusting the source name alone.")
    add_heading(doc, "To monitor calls for reviewers", 2)
    add_body(doc, "Use data/reviewer-call-sources.json. Add the official venue URL, allowlisted official domains, venue identity, role, and expected application route. Do not add a standalone form as the official source.")
    add_heading(doc, "To change labels and catalog behavior", 2)
    add_body(doc, "Use data/catalog-config.json for venue-type labels, source labels, ranking names, and CFP-link policy patterns.")
    add_heading(doc, "Before publishing a change", 2)
    for item in [
        "Run the relevant parser and policy tests.",
        "Run the production build.",
        "Open the local site and verify at least one affected card and link.",
        "Check that a past deadline disappears at the exact instant.",
        "Confirm that no unrelated user data or configuration was overwritten.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Repository map for a product builder", 1)
    add_table(doc,
        ["Area", "Main files", "System responsibility"],
        [
            ("Page shell", "app/page.jsx", "Chooses which tab is visible."),
            ("Main dashboard", "components/Dashboard.jsx", "Loads CFP data and applies user filters."),
            ("Cards and clocks", "components/CFPCard.jsx, Countdown.jsx, useDeadlineClock.js", "Displays one opportunity, its links, local time, and lifecycle."),
            ("CFP API", "app/api/cfps/route.js, lib/cfp.js", "Returns only current, deduplicated, correctly linked cards."),
            ("OpenReview", "lib/openreviewLive.js, lib/openreviewDirectory.js, scripts/scrape.mjs", "Discovers, hydrates, groups, snapshots, and merges OpenReview calls."),
            ("Link policy", "lib/cfpLinks.js", "Chooses the official CFP and prevents organizer-call cross-links."),
            ("Lifecycle", "lib/callLifecycle.js", "Decides whether a call is open at this exact moment."),
            ("Deduplication", "lib/dedupeCalls.js", "Prevents repeated cards without collapsing unrelated venues."),
            ("Hybrid search", "lib/hybridSearch.js", "Parses prompt constraints, expands selected concepts, and ranks current cards with BM25."),
            ("Trust", "lib/admissionPolicy.js, lib/legitimacy.js", "Screens candidates and records credibility evidence."),
            ("Other discovery", "scripts/discover-*.mjs, lib/webDiscovery.js", "Finds workshops, books, reviewer calls, publisher calls, and non-OpenReview sources."),
            ("Automation", ".github/workflows/pipeline.yml", "Runs the complete two-day workflow."),
            ("Configuration", "data/*-sources.json, watchlist.json, catalog-config.json", "Lets operators change monitored sources and labels."),
        ],
        [1600, 3300, 4460]
    )
    add_heading(doc, "Architecture takeaway", 1)
    add_body(doc, "The strongest part of CFP Radar is not any single scraper or AI model. It is the system design: discover broadly, trust narrowly, store evidence, separate opportunity types, remove expired records at the exact time, and keep the user-facing path fast and predictable.")
    add_body(doc, "At the current stage, this architecture is a sensible shared-catalog product. The next major architectural change should be driven by proven demand for separate lab ownership and personalization—not by user count alone.")

    doc.core_properties.title = "How CFP Radar Works"
    doc.core_properties.subject = "Product architecture and systems guide"
    doc.core_properties.author = "CFP Radar"
    doc.core_properties.keywords = "CFP Radar, architecture, product guide, OpenReview, academic calls"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
